#!/usr/bin/env python3
"""
Build the v5 (revision 1) submission files: manuscript DOCX, supplement DOCX,
and response-letter PDF, from the markdown sources in
notebooks/pearl/revision_1/.

Manuscript/supplement follow the same pandoc + python-docx pipeline as
build_pearl_docx.py (Cambria font, 1.5 line spacing, black text, bordered
tables, injected figure PNGs). The response letter is PDF per the journal's
stated requirement for the point-by-point response.
"""
import subprocess
import sys
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

PEARL_DIR = Path("/Users/sanjaybasu/waymark-local/notebooks/pearl")
REV_DIR = PEARL_DIR / "revision_1"
FIGS_DIR = PEARL_DIR / "figures_png"
OUTPUT_DIR = REV_DIR

DOCX_FILES = [
    {
        "src": REV_DIR / "manuscript" / "PEARL_manuscript_v5.md",
        "out": OUTPUT_DIR / "PEARL_manuscript_v5.docx",
        "figures": {
            "Figure 1.": FIGS_DIR / "fig1_pipeline.png",
            "Figure 2.": FIGS_DIR / "fig3_evaluator_dependence.png",
        },
    },
    {
        "src": REV_DIR / "supplement" / "PEARL_supplement_v5.md",
        "out": OUTPUT_DIR / "PEARL_supplement_v5.docx",
        "figures": {
            "Appendix Figure 1.": FIGS_DIR / "fig_patient_flow.png",
            "Appendix Figure 2.": FIGS_DIR / "fig4_camden.png",
        },
    },
]

PDF_FILE = {
    "src": REV_DIR / "response" / "response_to_reviewers.md",
    "out": REV_DIR / "response" / "response_to_reviewers.pdf",
}


def run_pandoc_docx(src: Path, out: Path) -> None:
    cmd = [
        "pandoc", str(src), "--to=docx",
        "--from=markdown+tex_math_dollars+pipe_tables+fenced_code_blocks+raw_html+superscript+subscript",
        "--standalone", "--output", str(out),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"  pandoc error:\n{result.stderr}", file=sys.stderr)
        sys.exit(1)


def get_para_text(para) -> str:
    return "".join(run.text for run in para.runs)


def set_all_text_black(doc: Document) -> None:
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)


def set_line_spacing_15(doc: Document) -> None:
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def set_font_cambria(doc: Document, font_name: str = "Cambria") -> None:
    try:
        doc.styles["Normal"].font.name = font_name
    except Exception:
        pass
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn("w:docDefaults"))
    if docDefaults is not None:
        rPrDefault = docDefaults.find(qn("w:rPrDefault"))
        if rPrDefault is None:
            rPrDefault = OxmlElement("w:rPrDefault")
            docDefaults.append(rPrDefault)
        rPr = rPrDefault.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            rPrDefault.append(rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:cs"), font_name)
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = font_name
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = font_name


def add_table_borders(doc: Document) -> None:
    border_sides = ["top", "left", "bottom", "right", "insideH", "insideV"]
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        old = tblPr.find(qn("w:tblBorders"))
        if old is not None:
            tblPr.remove(old)
        tblBorders = OxmlElement("w:tblBorders")
        for side in border_sides:
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            tblBorders.append(border)
        tblPr.append(tblBorders)


def insert_image_after_para(doc: Document, para_index: int, img_path: Path) -> None:
    target_para = doc.paragraphs[para_index]
    pic_para = doc.add_paragraph()
    run = pic_para.add_run()
    run.add_picture(str(img_path), width=Inches(6.0))
    body = doc.element.body
    new_para_xml = pic_para._element
    body.remove(new_para_xml)
    target_para._element.addnext(new_para_xml)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def inject_figures(doc: Document, figures: dict) -> None:
    if not figures:
        return
    injections = []
    for i, para in enumerate(doc.paragraphs):
        text = get_para_text(para).strip()
        for legend_start, fig_path in figures.items():
            if text.startswith(legend_start) and fig_path.exists():
                injections.append((i, fig_path, legend_start))
                break
    if not injections:
        print("  WARNING: No figure legend paragraphs found — figures not injected")
        return
    offset = 0
    for para_idx, fig_path, legend in injections:
        insert_image_after_para(doc, para_idx + offset, fig_path)
        offset += 1
        print(f"  Injected {fig_path.name} after '{legend}' paragraph")


def process_docx(spec: dict) -> None:
    src, out, figures = spec["src"], spec["out"], spec["figures"]
    print(f"\n{'='*60}\nProcessing: {src.name} -> {out.name}")
    run_pandoc_docx(src, out)
    doc = Document(str(out))
    set_font_cambria(doc)
    set_all_text_black(doc)
    set_line_spacing_15(doc)
    add_table_borders(doc)
    inject_figures(doc, figures)
    doc.save(str(out))
    print(f"  Saved: {out}  ({out.stat().st_size // 1024} KB)")


def build_response_pdf() -> None:
    src, out = PDF_FILE["src"], PDF_FILE["out"]
    print(f"\n{'='*60}\nProcessing: {src.name} -> {out.name}")
    header = Path("/tmp/pearl_response_header.tex")
    header.write_text(
        r"\usepackage{longtable}\usepackage{booktabs}\usepackage{array}"
        r"\renewcommand{\arraystretch}{1.2}"
        r"\usepackage{colortbl}"
    )
    cmd = [
        "pandoc", str(src),
        "--from=markdown+pipe_tables+fenced_code_blocks+raw_html+superscript+subscript",
        "--to=pdf", "--pdf-engine=xelatex", "--standalone",
        "--output", str(out),
        "-V", "geometry:margin=0.9in",
        "-V", "mainfont=Times New Roman",
        "-V", "fontsize=10pt",
        "-V", "linestretch=1.15",
        "--include-in-header", str(header),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"  pandoc error:\n{result.stderr[-4000:]}", file=sys.stderr)
        sys.exit(1)
    print(f"  Saved: {out}  ({out.stat().st_size // 1024} KB)")


def main():
    for spec in DOCX_FILES:
        if not spec["src"].exists():
            print(f"ERROR: source not found: {spec['src']}", file=sys.stderr)
            sys.exit(1)
    for spec in DOCX_FILES:
        process_docx(spec)
    build_response_pdf()
    print("\n" + "=" * 60)
    print("All v5 submission files built.")


if __name__ == "__main__":
    main()
