#!/usr/bin/env python3
"""
Build PEARL submission DOCX files from Markdown sources.
Steps per file:
  1. pandoc MD -> DOCX (OMML math = Cambria Math in Word; tables embedded)
  2. python-docx post-processing:
     a. Set all text runs to black (RGBColor 0,0,0)
     b. Apply 1.5 line spacing throughout
     c. Add full inside+outside borders to all tables
     d. Inject figure PNGs after their legend paragraphs
  3. Save final DOCX to notebooks/pearl/
"""
import subprocess
import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Inches, Pt, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING

PEARL_DIR = Path("/Users/sanjaybasu/waymark-local/notebooks/pearl")
FIGS_DIR = PEARL_DIR / "figures_png"
OUTPUT_DIR = PEARL_DIR

FILES = [
    {
        "src": PEARL_DIR / "manuscript" / "PEARL_manuscript_v4.md",
        "out": OUTPUT_DIR / "PEARL_manuscript_v4.docx",
        "figures": {
            "Figure 1.": FIGS_DIR / "fig_1_nano.png",
            "Figure 2.": FIGS_DIR / "fig2_drope_forest.png",
        },
    },
    {
        "src": PEARL_DIR / "supplement" / "PEARL_supplement_v4.md",
        "out": OUTPUT_DIR / "PEARL_supplement_v4.docx",
        "figures": {
            "Appendix Figure 1.": FIGS_DIR / "fig_patient_flow.png",
            "Appendix Figure 2.": FIGS_DIR / "fig3_imi_equity.png",
            "Appendix Figure 3.": FIGS_DIR / "fig4_camden.png",
            "Appendix Figure 4.": FIGS_DIR / "fig5_sensitivity.png",
        },
    },
    {
        "src": PEARL_DIR / "cover_letter" / "cover_letter_v4.md",
        "out": OUTPUT_DIR / "PEARL_cover_letter_v4.docx",
        "figures": {},
    },
]


def run_pandoc(src: Path, out: Path) -> None:
    cmd = [
        "pandoc",
        str(src),
        "--to=docx",
        "--from=markdown+tex_math_dollars+pipe_tables+fenced_code_blocks+raw_html+superscript+subscript",
        "--standalone",
        "--output", str(out),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"  pandoc error:\n{result.stderr}", file=sys.stderr)
        sys.exit(1)


def get_para_text(para) -> str:
    return "".join(run.text for run in para.runs)


def set_all_text_black(doc: Document) -> None:
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)


def set_line_spacing_15(doc: Document) -> None:
    """Apply 1.5 line spacing to every paragraph in the document."""
    for para in doc.paragraphs:
        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    # Also handle paragraphs inside table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    pf = para.paragraph_format
                    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def set_font_cambria(doc: Document, font_name: str = "Cambria") -> None:
    """Set the default body font to Cambria throughout the document."""
    # 1. Set Normal style font
    try:
        doc.styles["Normal"].font.name = font_name
    except Exception:
        pass

    # 2. Override document-level default fonts in XML
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn("w:docDefaults"))
    if docDefaults is not None:
        rPrDefault = docDefaults.find(qn("w:rPrDefault"))
        if rPrDefault is None:
            rPrDefault = OxmlElement("w:rPrDefault")
            docDefaults.append(rPrDefault)
        rPr = rPrDefault.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            rPrDefault.append(rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:cs"), font_name)

    # 3. Set font on every run
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = font_name
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = font_name


def _make_border_elm(val="single", sz="4", space="0", color="000000"):
    """Return an XML border element (w:top / w:bottom etc. — caller adds tag)."""
    elm = OxmlElement("w:top")  # tag overwritten by caller
    elm.set(qn("w:val"), val)
    elm.set(qn("w:sz"), sz)
    elm.set(qn("w:space"), space)
    elm.set(qn("w:color"), color)
    return elm


def add_table_borders(doc: Document) -> None:
    """Add single-line inside and outside borders to every table."""
    border_sides = ["top", "left", "bottom", "right", "insideH", "insideV"]

    for table in doc.tables:
        tbl = table._tbl
        # Find or create w:tblPr
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        # Remove old tblBorders if present
        old = tblPr.find(qn("w:tblBorders"))
        if old is not None:
            tblPr.remove(old)

        tblBorders = OxmlElement("w:tblBorders")
        for side in border_sides:
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")    # 0.5 pt
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            tblBorders.append(border)
        tblPr.append(tblBorders)


def insert_image_after_para(doc: Document, para_index: int, img_path: Path) -> None:
    target_para = doc.paragraphs[para_index]
    pic_para = doc.add_paragraph()
    run = pic_para.add_run()
    run.add_picture(str(img_path), width=Inches(6.0))
    body = doc.element.body
    new_para_xml = pic_para._element
    body.remove(new_para_xml)
    target_para._element.addnext(new_para_xml)
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Give the image paragraph 1.5 spacing too
    pic_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def inject_figures(doc: Document, figures: dict) -> None:
    if not figures:
        return
    injections = []
    for i, para in enumerate(doc.paragraphs):
        text = get_para_text(para).strip()
        for legend_start, fig_path in figures.items():
            if text.startswith(legend_start) and fig_path.exists():
                injections.append((i, fig_path, legend_start))
                break
    if not injections:
        print("  WARNING: No figure legend paragraphs found — figures not injected")
        return
    offset = 0
    for i, (para_idx, fig_path, legend) in enumerate(injections):
        adjusted_idx = para_idx + offset
        insert_image_after_para(doc, adjusted_idx, fig_path)
        offset += 1
        print(f"  Injected {fig_path.name} after '{legend}' paragraph")


def process_file(spec: dict) -> None:
    src = spec["src"]
    out = spec["out"]
    figures = spec["figures"]

    print(f"\n{'='*60}")
    print(f"Processing: {src.name} -> {out.name}")

    print(f"  Running pandoc...")
    run_pandoc(src, out)
    print(f"  pandoc done: {out}")

    print(f"  Post-processing with python-docx...")
    doc = Document(str(out))

    set_font_cambria(doc)
    print(f"  Font set to Cambria")

    set_all_text_black(doc)
    print(f"  All text set to black")

    set_line_spacing_15(doc)
    print(f"  1.5 line spacing applied")

    add_table_borders(doc)
    print(f"  Table borders applied (inside + outside)")

    inject_figures(doc, figures)

    doc.save(str(out))
    print(f"  Saved: {out}")
    size_kb = out.stat().st_size // 1024
    print(f"  File size: {size_kb} KB")


def main():
    print("PEARL DOCX Build Pipeline")
    print(f"Output dir: {OUTPUT_DIR}")

    for spec in FILES:
        if not spec["src"].exists():
            print(f"ERROR: source not found: {spec['src']}", file=sys.stderr)
            sys.exit(1)

    for spec in FILES:
        process_file(spec)

    print("\n" + "="*60)
    print("All DOCX files built successfully.")
    print("\nOutput files:")
    for spec in FILES:
        out = spec["out"]
        if out.exists():
            print(f"  {out}  ({out.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
